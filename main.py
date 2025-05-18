import os
import tempfile
from typing import List
from docx import Document
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
from Model.question_output import Answer, Question
from helper.question_modification import classify_line, classify_line_ml, question_modify , dataframe_convert
from AI.train_model import learning_set

import re
app = FastAPI()

@app.post("/extract-questions")
async def extract_questions(file: UploadFile = File(...)):
    # Lưu tạm file docx
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        # Đọc nội dung từ docx
        doc = Document(tmp_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip() != ""]
        classified = [(line, classify_line_ml(line)) for line in paragraphs]

            # Gom nhóm thành câu hỏi
        questions_raw = question_modify(classified)
        qs : list[Question] = []
        for i in range(len(questions_raw)):
            if any(questions_raw[i].get(k) in [None, []] for k in ["correct", "answers", "question"]):
                continue    
            try:
                qs.append(
                    Question(
                        " ".join(questions_raw[i]["question"]),
                        [Answer(a) for a in questions_raw[i]["answers"]],
                        questions_raw[i]["correct"]
                        )
                    )
            except Exception as e:
                print(f"Error: {e}")
                print(f"Question: {questions_raw[i]['question']}")
                print(f"Answers: {questions_raw[i]['answers']}")
                print(f"Correct: {questions_raw[i]['correct']}")
                continue
        
        # Chuyển sang JSON-friendly
        result = []
        for q in qs:
            result.append(q.json_convert())
        return JSONResponse(content=result)

    finally:
        os.remove(tmp_path)


if __name__ == "__main__":
    ...
    # --- Sử dụng hàm:
    # file_path = "questions/cauhoi.docx"  

    # doc = Document(file_path)

    # paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip() != ""]
    # # classified = [(line, classify_line(line)) for line in paragraphs] 
    # # [1082:1093] [1524:1530]
    # classified = []
    # previous_type = ""
    # for line in paragraphs:
        
    #     current_line = (line,classify_line_ml(line))
        
    #     # if current_line[1] == "answer" and not re.match(r"^[a-dA-D]\.", line):
    #     #     state.CURRENT_RETURN_TYPE = (previous_type,0)
    #     #     current_line = (line, classify_line(line))
            
    #     classified.append(current_line)
    #     previous_type = current_line[1]
    # # count = 0
    # # for text, label in classified:
    # #     count += 1
    # #     print(f"{count} [{label.upper()}] {text}")
    
    
    # questions = question_modify(classified)
    # # dataframe_convert(questions)
    # # learning_set()
    # qs : list[Question] = []
    # for i in range(len(questions)):
    #     if questions[i]["correct"] is None or questions[i]["answers"] is None or questions[i]["question"] is None:
    #         continue
    #     try:
    #         qs.append(
    #             Question(
    #                 " ".join(questions[i]["question"]),
    #                 [Answer(a) for a in questions[i]["answers"]],
    #                 questions[i]["correct"]
    #                 )
    #             )
    #     except Exception as e:
    #         print(f"Error: {e}")
    #         print(f"Question: {questions[i]['question']}")
    #         print(f"Answers: {questions[i]['answers']}")
    #         print(f"Correct: {questions[i]['correct']}")
    #         continue
    # # for q in questions:
    # #     print("\n--- Câu hỏi ---")
    # #     print(q["question"])
    # #     print("Các đáp án:")
    # #     for a in q["answers"]:
    # #         print(f" - {a}")
    # #     print(f"{q["correct"]}")
    # for q in qs:
    #     print(q)    
    # print("✅ Đã chuyển đổi xong.")

